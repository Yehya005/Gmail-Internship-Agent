"""Convert report.md → report.docx with sensible Word styling.

Targets the IEEE-paper conventions where this is reasonable in a single
.docx (Title centered, two-column body, Times-New-Roman 10pt). The
final layout-pass into the official IEEE template is still a copy-
paste job, but the output here is a clean Word document a reviewer can
open and read directly.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

SRC = Path("report.md")
DST = Path("report.docx")


def _set_cols(section, n_cols: int = 2) -> None:
    """Switch a section to N columns. python-docx exposes this only via
    raw XML, so we add a `<w:cols>` element to the section properties."""
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(n_cols))
    cols.set(qn("w:space"), "360")  # 0.25 inch gutter


def _add_paragraph_with_inline(doc, line: str, *, style: str = "Body Text") -> None:
    """Add a paragraph rendering inline **bold**, *italic*, and `code`."""
    p = doc.add_paragraph(style=style)
    pos = 0
    pat = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    for m in pat.finditer(line):
        if m.start() > pos:
            p.add_run(line[pos:m.start()])
        tok = m.group(0)
        run = p.add_run(tok[2:-2] if tok.startswith("**")
                        else tok[1:-1])
        if tok.startswith("**"):
            run.bold = True
        elif tok.startswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run.italic = True
        pos = m.end()
    if pos < len(line):
        p.add_run(line[pos:])


def main() -> int:
    if not SRC.exists():
        print(f"error: {SRC} not found", file=sys.stderr)
        return 1
    text = SRC.read_text(encoding="utf-8")

    doc = Document()

    # Page margins — IEEE uses 0.75" left/right, 1" top/bottom roughly.
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # Default body style — Times New Roman 10pt as in the IEEE template.
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    # Walk the markdown line by line. The report uses a small subset of
    # markdown features so a hand-rolled parser is enough.
    lines = text.split("\n")
    i = 0
    in_code = False
    code_buffer: list[str] = []
    title_done = False
    columns_started = False

    def flush_code():
        nonlocal code_buffer
        if not code_buffer:
            return
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.1)
        run = para.add_run("\n".join(code_buffer))
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        code_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Code fences — toggle code accumulation
        if stripped.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Section dividers
        if stripped == "---":
            i += 1
            continue

        # Title: first H1 line
        if stripped.startswith("# ") and not title_done:
            title = stripped[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(16)
            title_done = True
            i += 1
            continue

        # Author / affiliation lines come right after the title before the
        # first ## section. Render them centered, smaller, italic-ish.
        if title_done and not columns_started and stripped.startswith("**"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Strip any markdown bold so we get clean text + add as italic
            text_only = re.sub(r"\*\*", "", stripped)
            run = p.add_run(text_only)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Once we hit the first H2 ("## Abstract"), switch to two-column
        if stripped.startswith("## ") and not columns_started:
            new_section = doc.add_section(WD_SECTION.CONTINUOUS)
            new_section.top_margin = Inches(0.75)
            new_section.bottom_margin = Inches(0.75)
            new_section.left_margin = Inches(0.6)
            new_section.right_margin = Inches(0.6)
            _set_cols(new_section, 2)
            columns_started = True

        # H2 (top-level body section)
        if stripped.startswith("## "):
            heading = stripped[3:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.bold = True
            run.font.size = Pt(11)
            i += 1
            continue

        # H3 (subsection)
        if stripped.startswith("### "):
            heading = stripped[4:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.bold = True
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # H4 (smaller subsection)
        if stripped.startswith("#### "):
            heading = stripped[5:].strip()
            p = doc.add_paragraph()
            run = p.add_run(heading)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Bulleted lists
        if stripped.startswith("- ") or stripped.startswith("* "):
            content = stripped[2:].strip()
            _add_paragraph_with_inline(doc, content, style="List Bullet")
            i += 1
            continue

        # Numbered lists
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            content = m.group(2)
            _add_paragraph_with_inline(doc, content, style="List Number")
            i += 1
            continue

        # Blank line — paragraph break
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        _add_paragraph_with_inline(doc, line)
        i += 1

    # If the markdown ended inside a code block (shouldn't happen) flush.
    flush_code()

    doc.save(DST)
    print(f"wrote {DST.resolve()}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
